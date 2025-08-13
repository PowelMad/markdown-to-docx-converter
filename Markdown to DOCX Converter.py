import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class MarkdownToDocxConverter:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self.current_list_level = 0
        self.in_code_block = False
        self.in_table = False
        self.table_data = []
        self.code_block_lines = []

    def _setup_styles(self):
        #"""Configuration des styles de base"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
        # Création de styles personnalisés
        styles = {
            'Heading1': {'type': 1, 'size': 16, 'bold': True, 'color': RGBColor(0, 0, 0)},
            'Heading2': {'type': 1, 'size': 14, 'color': RGBColor(0, 51, 102)},
            'Heading3': {'type': 1, 'size': 12, 'color': RGBColor(0, 102, 102)},
            'Heading4': {'type': 1, 'size': 11, 'italic': True},
            'CodeChar': {'type': 2, 'name': 'Consolas', 'size': 10, 'color': RGBColor(0, 0, 0)},
            'QuoteChar': {'type': 2, 'italic': True, 'color': RGBColor(102, 102, 102)}
        }
        
        for style_name, attrs in styles.items():
            if style_name not in self.doc.styles:
                new_style = self.doc.styles.add_style(style_name, attrs['type'])
                for attr, value in attrs.items():
                    if attr == 'type':
                        continue
                    elif attr == 'color':
                        new_style.font.color.rgb = value
                    elif attr == 'name':
                        new_style.font.name = value
                    else:
                        setattr(new_style.font, attr, value)
    def convert(self, md_file_path, docx_file_path):
        """Convertit un fichier Markdown en document Word."""
        if not os.path.exists(md_file_path):
            raise FileNotFoundError(f"Le fichier {md_file_path} n'existe pas")
        
        try:
            with open(md_file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
        except UnicodeDecodeError:
            with open(md_file_path, 'r', encoding='latin-1') as f:
                md_content = f.read()

        lines = md_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if self.in_code_block:
                self._process_code_block(line, i, lines)
                i += 1
                continue
                
            if self.in_table:
                self._process_table(line, i, lines)
                i += 1
                continue

            self._process_line(line)
            i += 1

        self._save_document(docx_file_path)

    def _process_line(self, line):
        """Traite une ligne de Markdown."""
        if line.startswith('# '):
            self._add_heading(line[2:], 0)
        elif line.startswith('## '):
            self._add_heading(line[3:], 1)
        elif line.startswith('### '):
            self._add_heading(line[4:], 2)
        elif line.startswith('#### '):
            self._add_heading(line[5:], 3)
        elif line.startswith('```'):
            self.in_code_block = True
        elif line.startswith('|') and line.count('|') >= 2:
            self.in_table = True
            self.table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
        elif line.startswith('>'):
            self._add_quote(line[1:].strip())
        elif line.startswith(('- ', '* ', '1. ', '□ ', '✅ ', '❌ ')):
            self._add_list_item(line)
        elif not line:
            self.doc.add_paragraph()
        else:
            self._add_formatted_paragraph(line)

    def _process_code_block(self, line, current_idx, lines):
        """Traite un bloc de code."""
        if line.startswith('```'):
            # Fin du bloc de code
            code_content = '\n'.join(self.code_block_lines)
            self._add_code_block(code_content)
            self.code_block_lines = []
            self.in_code_block = False
        else:
            self.code_block_lines.append(line)

    def _process_table(self, line, current_idx, lines):
        """Traite un tableau."""
        if '|' in line and line.count('|') >= 2:
            self.table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
        else:
            # Fin du tableau
            self._add_table(self.table_data)
            self.table_data = []
            self.in_table = False
            # Retourner en arrière pour traiter la ligne actuelle
            self._process_line(line.strip())

    def _add_heading(self, text, level):
        """Ajoute un titre."""
        heading = self.doc.add_heading(text, level=level)
        style_map = {
            0: 'Heading1',
            1: 'Heading2', 
            2: 'Heading3',
            3: 'Heading4'
        }
        heading.style = self.doc.styles[style_map[level]]

    def _add_list_item(self, text):
        """Ajoute un élément de liste."""
        if text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            list_style = 'List Number'
        else:
            list_style = 'List Bullet'
        
        p = self.doc.add_paragraph(style=list_style)
        p.paragraph_format.left_indent = Pt(18)
        
        # Gestion des cases à cocher
        if text.startswith('✅'):
            text = "✓ " + text[1:].strip()
        elif text.startswith('❌'):
            text = "✗ " + text[1:].strip()
        
        clean_text = text.lstrip('-*123456789. ')
        self._add_formatted_run(p, clean_text)

    def _add_table(self, data):
        """Ajoute un tableau."""
        if not data:
            return
            
        # Normaliser le nombre de colonnes
        max_cols = max(len(row) for row in data)
        for row in data:
            while len(row) < max_cols:
                row.append('')
        # Créer le tableau avec le nombre de lignes et de colonnes
        table = self.doc.add_table(rows=len(data), cols=max_cols)
        table.style = 'Light Shading Accent 1'
        # Remplir le tableau avec les données
        for row_idx, row in enumerate(data):
            for col_idx, cell in enumerate(row):
                if col_idx < len(table.columns):
                    table.cell(row_idx, col_idx).text = str(cell)
                    if row_idx == 0:
                        table.cell(row_idx, col_idx).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_code_block(self, code):
        """Ajoute un bloc de code."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(code)
        run.style = self.doc.styles['CodeChar']  # Changé de 'Code' à 'CodeChar'
        
        try:
            shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
            p._p.get_or_add_pPr().append(shading_elm)
        except Exception:
            pass

    def _add_quote(self, text):
        """Ajoute une citation."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(text)
        run.style = self.doc.styles['QuoteChar']  # Changé de 'Quote' à 'QuoteChar'
        
        try:
            border_xml = r'<w:pBdr {}><w:left w:val="single" w:sz="4" w:space="9" w:color="CCCCCC"/></w:pBdr>'.format(nsdecls('w'))
            p._p.get_or_add_pPr().append(parse_xml(border_xml))
        except Exception:
            pass
    def _add_formatted_paragraph(self, text):
            """Ajoute un paragraphe formaté."""
            p = self.doc.add_paragraph()
            self._add_formatted_run(p, text)

    def _add_formatted_run(self, paragraph, text):
        """Ajoute du texte formaté à un paragraphe."""
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.style = self.doc.styles['CodeChar']  # Changé ici
            else:
                run = paragraph.add_run(part)
                self._apply_emoji_formatting(run, part)

    def _apply_emoji_formatting(self, run, text):
        """Applique le formatage des emojis."""
        emoji_colors = {
            '🚀': RGBColor(255, 102, 0),  # Orange
            '⚠️': RGBColor(255, 0, 0),    # Rouge
            '💡': RGBColor(0, 102, 204),  # Bleu
            '🎯': RGBColor(0, 153, 0)     # Vert
        }
        
        for emoji, color in emoji_colors.items():
            if emoji in text:
                run.font.color.rgb = color
                break
    # Save the document to the specified file path.
    def _save_document(self, file_path):
        """Enregistre le document Word."""
        try:
            self.doc.save(file_path)
            print(f"✅ Conversion réussie : {file_path}")
            return True
        except PermissionError:
            print(f"❌ Erreur : Impossible d'écrire dans {file_path} (permission refusée)")
        except Exception as e:
            print(f"❌ Erreur inattendue lors de l'enregistrement : {e}")
        return False

def main():
    converter = MarkdownToDocxConverter()
    # Chemin du fichier Markdown et du fichier de sortie
    input_file = 'ebook_aeo_complete (2).md'#Nom du fichier Markdown à convertir
    output_file = 'AEO_Playbook.docx'#Nom du fichier de sortie Word

    # Conversion du fichier Markdown en Word
    print(f"🔄 Conversion de {input_file} vers {output_file}...")

    # Essayer de convertir le fichier
    try:
        converter.convert(input_file, output_file)
    except FileNotFoundError as e:
        print(f"❌ Erreur : {e}")
    except Exception as e:
        print(f"❌ Erreur inattendue : {e}")

if __name__ == "__main__":
    main()